# -*- coding: utf-8 -*-
"""
Generate comprehensive Word documents (Technical Report and Demo Script)
for the TerriMap commercial territory design project.
Version 3.5: Extremely detailed academic & source code walkthrough.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUTPUT_DIR = os.path.expanduser("~\\Desktop")

# ============================================================
# HELPERS FOR WORD FORMATTING
# ============================================================

def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elem)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set cell padding in twentieths of a point (dxa)."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = tcPr.makeelement(qn('w:tcMar'), {})
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = tcMar.makeelement(qn(f'w:{side}'), {
            qn('w:w'): str(val),
            qn('w:type'): 'dxa'
        })
        tcMar.append(node)
    tcPr.append(tcMar)

def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a beautifully formatted table with zebra striping and custom padding."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Enable borders
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = tblPr[0].makeelement(qn('w:tblBorders'), {})
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = borders.makeelement(qn(f'w:{border_name}'), {
                qn('w:val'): 'single',
                qn('w:sz'): '4',
                qn('w:space'): '0',
                qn('w:color'): 'D3D3D3'
            })
            borders.append(border)
        tblPr[0].append(borders)

    # Format Header Row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "1B365D")  # Deep Navy Blue
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Format Data Rows
    for r_idx, row in enumerate(rows):
        shading_color = "F7F9FB" if r_idx % 2 == 1 else "FFFFFF"  # Zebra striping
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            set_cell_shading(cell, shading_color)
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    return table

def add_formula(doc, text, bold=False):
    """Add a mathematical formula block styled like a callout."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x2E, 0x40, 0x53)
    run.bold = bold
    return p

def add_code_block(doc, code_lines):
    """Add a formatted code block with monospace font and gray border background."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F2F4F4")  # Light gray background
    set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
    
    # Left border color accent
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = tcPr.makeelement(qn('w:tcBorders'), {})
    left_border = tcBorders.makeelement(qn('w:left'), {
        qn('w:val'): 'single',
        qn('w:sz'): '24', # 3pt thick
        qn('w:space'): '0',
        qn('w:color'): '1B365D' # Navy accent
    })
    tcBorders.append(left_border)
    # Remove others
    for border_name in ['top', 'bottom', 'right']:
        b = tcBorders.makeelement(qn(f'w:{border_name}'), {qn('w:val'): 'none'})
        tcBorders.append(b)
    tcPr.append(tcBorders)

    # Set content
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    
    run = p.add_run("\n".join(code_lines))
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x23, 0x2B, 0x2B)

def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Primary color

def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(0x4A, 0x77, 0x9D) # Secondary color

def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11.5)
    run.bold = True
    run.italic = True
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def para(doc, text, bold=False, italic=False, space_after=6, space_before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Arial'
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.75 * (level + 1))
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_callout(doc, text, title="CHÚ Ý"):
    """Adds a beautiful warning/info callout block."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "FEF9E7")  # Very light warm yellow
    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
    
    # Border
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = tcPr.makeelement(qn('w:tcBorders'), {})
    left_border = tcBorders.makeelement(qn('w:left'), {
        qn('w:val'): 'single',
        qn('w:sz'): '24',
        qn('w:space'): '0',
        qn('w:color'): 'D68910' # Orange accent
    })
    tcBorders.append(left_border)
    for b_name in ['top', 'bottom', 'right']:
        tcBorders.append(tcBorders.makeelement(qn(f'w:{b_name}'), {qn('w:val'): 'none'}))
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    
    run_title = p.add_run(f"★ {title}: ")
    run_title.bold = True
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(10)
    run_title.font.color.rgb = RGBColor(0xB7, 0x79, 0x1F)
    
    run_text = p.add_run(text)
    run_text.italic = True
    run_text.font.name = 'Arial'
    run_text.font.size = Pt(10)
    run_text.font.color.rgb = RGBColor(0x5F, 0x4B, 0x1A)

# ============================================================
# GENERATION: FILE 1 - TECHNICAL REPORT
# ============================================================

def build_report():
    doc = Document()
    
    # Document global styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Page Setup - Standards
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)
    
    # --- TITLE PAGE ---
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(120)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run(
        'BÁO CÁO KHẢO SÁT & THIẾT KẾ KỸ THUẬT CHUYÊN SÂU\n'
        'HỆ THỐNG PHÂN CHIA VÙNG THƯƠNG MẠI THỜI GIAN THỰC (TERRIMAP)\n'
        'COMMERCIAL TERRITORY DESIGN SYSTEM'
    )
    run_title.font.size = Pt(18)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(180)
    run_sub = p_sub.add_run(
        'TerriMap Enterprise System (Version 3.5)\n'
        'Tài liệu nghiên cứu chuyên sâu về Kiến trúc, Giải thuật, UI/UX và Database\n'
        'Phục vụ công tác bảo vệ đồ án tốt nghiệp ngành Công nghệ thông tin / Khoa học máy tính'
    )
    run_sub.font.size = Pt(12)
    run_sub.italic = True
    run_sub.font.color.rgb = RGBColor(0x4A, 0x77, 0x9D)
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_after = Pt(20)
    run_meta = p_meta.add_run(
        'Tác giả chuyên môn: thiendominh0-star (Senior Technical Architect)\n'
        'Phạm vi nghiên cứu: Hình học tính toán · Metaheuristics · Lập trình song song trình duyệt · RBAC & RLS Security\n'
        'Công nghệ triển khai: React 18 · TypeScript · Vite · Zustand · Supabase PostgreSQL · Leaflet Maps\n'
        'Cấp tài liệu: Mật - Tài liệu kỹ thuật chi tiết nhất'
    )
    run_meta.font.size = Pt(9.5)
    run_meta.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    doc.add_page_break()
    
    # --- TABLE OF CONTENTS ---
    h1(doc, 'MỤC LỤC CHI TIẾT BÁO CÁO')
    toc_items = [
        '1. Tổng quan hệ thống và bài toán nghiệp vụ',
        '   1.1. Bối cảnh và bài toán phân chia vùng thương mại',
        '   1.2. Các ràng buộc nghiệp vụ phi tuyến tính',
        '   1.3. Thống kê quy mô và chỉ số hệ thống thực tế',
        '2. Kiến trúc phần mềm chi tiết (5-Layer Clean Architecture)',
        '   2.1. Phân tích vai trò từng tầng liên kết một chiều',
        '   2.2. Luồng dữ liệu (Data Flow) và cơ chế Facade pattern',
        '3. Phân tích chi tiết Zustand State Management và Offline-First',
        '   3.1. Thiết kế Zustand store trong dataStore.ts',
        '   3.2. Cơ chế Optimistic Update và an toàn luồng lưu dữ liệu (Saving indicator)',
        '   3.3. Thuật toán Offline-First LocalStorage và đồng bộ Supabase',
        '4. Phân hệ hiển thị UI/UX và Bản đồ tương tác Leaflet (L4 - Presentation)',
        '   4.1. Thư viện hiển thị được sử dụng (Leaflet, Recharts, Tailwind)',
        '   4.2. Quản lý trạng thái UI Store (uiStore.ts) chuyên biệt',
        '   4.3. Thành phần TerritoryMap.tsx và cơ chế MapFlyTo nâng cao',
        '   4.4. Quy tắc tô màu trực quan và hiển thị chỉ báo lỗi địa lý (Island, Disconnected)',
        '   4.5. Phân tích mã nguồn và chức năng của 5 màn hình chính',
        '5. Phân tích chi tiết thuật toán phân vùng và liên quan đến code',
        '   5.1. Lựa chọn hạt nhân tối ưu (Farthest-Point Seeding Heuristic)',
        '   5.2. Thuật toán Greedy Seed Expansion và chiến lược liên thông Grow-to-Reach',
        '   5.3. Thuật toán Local Search (2-opt Improvement) kiểm soát biên liên thông',
        '   5.4. Thuật toán Simulated Annealing (Ủ mô phỏng) & Tối ưu hóa phân phối Boltzmann',
        '   5.5. Song song hóa luồng phụ qua Web Workers giải quyết nghẽn Main Thread',
        '6. Các công thức toán học và hình học cốt lõi',
        '   6.1. Công thức khoảng cách mặt cầu Haversine và chống tràn số thực',
        '   6.2. Công thức tâm hình học đa giác Shoelace và phòng vệ đa giác suy biến',
        '   6.3. Xây dựng đồ thị liên kết ma trận kề (Adjacency Matrix)',
        '7. Thiết kế cơ sở dữ liệu và Cơ chế bảo mật phân quyền (Supabase RLS & RBAC)',
        '   7.1. Lược đồ quan hệ thực thể cơ sở dữ liệu (Database Schema)',
        '   7.2. Chính sách Row Level Security (RLS) bảo mật đa dự án',
        '   7.3. Triggers tự động hóa đồng bộ tài khoản người dùng',
        '8. Bộ 30 câu hỏi phản biện chuyên sâu và Lời giải thích học thuật bảo vệ đồ án',
        '9. Đánh giá chất lượng hệ thống và định hướng mở rộng tương lai'
    ]
    for item in toc_items:
        if item.startswith('   '):
            bullet(doc, item.strip())
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.space_before = Pt(8)
            run = p.add_run(item)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
            
    doc.add_page_break()
    
    # ===================== PHẦN 1 =====================
    h1(doc, '1. TỔNG QUAN HỆ THỐNG VÀ BÀI TÓAN NGHIỆP VỤ')
    
    h2(doc, '1.1. Bối cảnh và bài toán phân chia vùng thương mại')
    para(doc, 'Trong quản trị phân phối thương mại hiện đại, việc phân chia thị trường bán lẻ cho lực lượng bán hàng (Sales Force) được gọi là bài toán Thiết kế vùng thương mại (Commercial Territory Design). Giả sử doanh nghiệp có một khu vực hoạt động lớn gồm n vùng địa lý nhỏ cơ bản (được gọi là các Zones, ví dụ: các Phường thuộc thành phố Hà Nội) và m nhân viên kinh doanh (Sales Agents). Mục tiêu cần đạt được là gom n zones này thành m cụm phân vùng lớn hơn (được gọi là các Districts), sao cho mỗi nhân viên phụ trách một cụm duy nhất.')
    para(doc, 'Tuy nhiên, đây không đơn giản chỉ là việc chia đều số lượng zone trên giấy tờ, mà nó đòi hỏi một sự tối ưu hóa toán học phức tạp nhằm giải quyết các mâu thuẫn lớn trong thực tiễn nghiệp vụ:')
    bullet(doc, 'Cân bằng tải công việc (Workload Balance): Phân chia sao cho lượng khách hàng (Customers), số lượng đơn hàng (Orders) hoặc doanh số (Revenue) giữa các nhân viên không lệch quá nhiều. Sự chênh lệch quá lớn sẽ dẫn đến sự mất công bằng trong chính sách doanh số, gây ức chế tâm lý hoặc quá tải công việc.')
    bullet(doc, 'Tối ưu hóa không gian (Compactness): Các zone trong cùng một cụm phải gọn gàng, giảm thiểu khoảng cách từ tâm cụm đến các điểm thành viên. Điều này giúp giảm bán kính di chuyển, tiết kiệm chi phí xăng xe và thời gian di chuyển của nhân viên thị trường.')
    bullet(doc, 'Tính liên thông địa lý cứng (Geographical Contiguity): Đây là ràng buộc bắt buộc. Nhân viên kinh doanh phải có khả năng di chuyển liên tục từ zone này qua zone khác trong khu vực của mình mà không cần cắt qua lãnh thổ của nhân viên khác. Một cụm bị đứt rời địa lý (disconnected) sẽ làm tăng quãng đường di chuyển và gây chồng chéo quản lý.')

    h2(doc, '1.2. Các ràng buộc nghiệp vụ phi tuyến tính')
    para(doc, 'Bài toán thiết kế vùng thương mại thuộc lớp bài toán tối ưu hóa đa mục tiêu với ràng buộc cứng (Constrained Multi-Objective Optimization). Về mặt lý thuyết, việc tìm lời giải tối ưu toàn cục là bài toán NP-hard. Số lượng phương án phân hoạch n tập hợp thành m cụm tăng cực kỳ nhanh theo số Stirling loại hai:')
    add_formula(doc, "S(n, m) = \\frac{1}{m!} \\sum_{i=0}^{m} (-1)^{m-i} \\binom{m}{i} i^n")
    para(doc, 'Ví dụ, với n = 50 zones và m = 5 districts, số lượng phương án phân chia khả thi là khoảng 10^32 phương án. Việc duyệt cạn (brute-force) là không khả thi trên bất kỳ máy tính siêu cấp nào. Do đó, TerriMap thiết kế một hệ thống Heuristic/Metaheuristic thông minh, kết hợp thuật toán Greedy, Local Search và Simulated Annealing để giải bài toán này chỉ trong vòng vài mili giây.')

    h2(doc, '1.3. Thống kê quy mô và chỉ số hệ thống thực tế')
    add_styled_table(doc,
        ['Tham số hệ thống', 'Mô tả chi tiết', 'Tầm quan trọng kỹ thuật'],
        [
            ['Mô hình hóa dữ liệu', '12 Zones mẫu phân bổ tại 3 vùng địa lý chính (Hà Nội, Huế, TP. HCM).', 'Làm dữ liệu cơ sở để thuật toán loang địa hình.'],
            ['Trọng số hoạt động', 'Hỗ trợ CUSTOMER, ORDER, REVENUE cho từng Zone.', 'Quyết định việc cân bằng tải đa mục tiêu qua hàm cost.'],
            ['Chỉ số kiểm thử tự động', '371 unit & integration tests chạy trên Vitest.', 'Đảm bảo không phát sinh lỗi logic toán học khi refactor code.'],
            ['Mức độ chính xác tính toán', 'Haversine WGS-84 với sai số khoảng cách dưới 0.1%.', 'Đảm bảo khoảng cách địa lý phản ánh đúng thực tế mặt cầu.'],
            ['Khả năng xử lý đa luồng', 'Tích hợp Web Workers chạy song song.', 'Giải phóng Main Thread của trình duyệt khỏi tính toán nặng.']
        ])

    # ===================== PHẦN 2 =====================
    doc.add_page_break()
    h1(doc, '2. KIẾN TRÚC PHẦN MỀM CHI TIẾT (5-LAYER CLEAN ARCHITECTURE)')
    
    h2(doc, '2.1. Phân tích vai trò từng tầng liên kết một chiều')
    para(doc, 'Ứng dụng TerriMap được thiết kế tuân thủ nghiêm ngặt mô hình 5-Layer Clean Architecture. Nguyên tắc cốt lõi của kiến trúc này là sự phụ thuộc một chiều hướng tâm (Dependency Rule): Các tầng bên ngoài phụ thuộc vào tầng bên trong, nhưng các tầng bên trong hoàn toàn không biết gì về các tầng bên ngoài.')
    
    add_styled_table(doc,
        ['Tầng kiến trúc', 'Thư mục & Tệp tin', 'Nhiệm vụ chi tiết', 'Quy tắc phụ thuộc'],
        [
            ['L4 — UI Layer', 'src/components/, src/pages/', 'Render giao diện React, vẽ bản đồ Leaflet, biểu đồ Recharts, xử lý sự kiện DOM.', 'Được phép import L3, L2, L1, L0.'],
            ['L3 — Facades Layer', 'src/context/FacadeContext.tsx', 'Cung cấp API tối giản cho UI sử dụng, che giấu sự phức tạp của dịch vụ database và các stores.', 'Được phép import L2, L1, L0. Nghiêm cấm chứa code UI.'],
            ['L2 — Services Layer', 'src/services/db.ts, src/store/dataStore.ts', 'Xử lý các nghiệp vụ (Business Rules), điều phối dữ liệu với DB Supabase, quản lý state.', 'Được phép import L1, L0. Không import UI.'],
            ['L1 — Libraries Layer', 'lib/geometry.ts, lib/partition.ts', 'Chứa các pure functions tính toán hình học, khoảng cách, và các thuật toán lõi.', 'Được phép import L0. Không side effects, không import framework.'],
            ['L0 — Domain Types', 'src/types/domain.ts', 'Định nghĩa Single Source of Truth cho các kiểu dữ liệu thực thể (Entity Types).', 'Không import bất cứ gì ngoài chính nó (độc lập tuyệt đối).']
        ])
    
    h2(doc, '2.2. Luồng dữ liệu (Data Flow) và cơ chế Facade pattern')
    para(doc, 'Cơ chế Facade pattern (L3) thông qua `FacadeContext.tsx` đóng vai trò quan trọng trong việc giảm độ kết nối (coupling) giữa UI và giải thuật. Khi người dùng click nút "Chạy thuật toán Simulated Annealing" trên màn hình:')
    bullet(doc, 'UI (L4) gọi hàm `runPartitioning(algoName, options)` thông qua Facade Context.')
    bullet(doc, 'Facade (L3) tiếp nhận yêu cầu, trích xuất dữ liệu Zones và Districts hiện tại từ Zustand Store (L2).')
    bullet(doc, 'Facade khởi tạo một Web Worker mới, gửi dữ liệu xuống Worker để Worker gọi hàm `partitionSimulatedAnnealing()` thuộc Library (L1) tính toán.')
    bullet(doc, 'Sau khi Worker tính xong, nó trả kết quả về Facade, Facade cập nhật Zustand Store, store lưu dữ liệu xuống Supabase DB (L2) ngầm dưới nền, và đồng thời cập nhật UI để re-render bản đồ (L4).')

    # ===================== PHẦN 3 =====================
    doc.add_page_break()
    h1(doc, '3. PHÂN TÍCH CHI TIẾT ZUSTAND STATE MANAGEMENT VÀ OFFLINE-FIRST')
    
    h2(doc, '3.1. Thiết kế Zustand store trong dataStore.ts')
    para(doc, 'Quản lý trạng thái dữ liệu nghiệp vụ của TerriMap được tập trung hóa tại `src/store/dataStore.ts`. Đây là nơi duy nhất giữ chân dữ liệu thực tế từ Database và điều phối các hành động ghi/đọc.')
    para(doc, 'Mã nguồn cấu trúc dữ liệu của Zustand dataStore bao gồm:')
    add_code_block(doc, [
        "interface DataState {",
        "  projects: Project[];",
        "  currentProject: Project | null;",
        "  regions: Region[];",
        "  zones: Zone[];",
        "  salesAgents: SalesAgent[];",
        "  assignments: Assignment[];",
        "  initialized: boolean;",
        "  saving: boolean;",
        "  error: string | null;",
        "  init: (projectId: string) => Promise<void>;",
        "  addZone: (zone: Omit<Zone, 'id'>) => Promise<void>;",
        "  updateAssignments: (newAssignments: Assignment[]) => Promise<void>;",
        "}"
    ])

    h2(doc, '3.2. Cơ chế Optimistic Update và an toàn luồng lưu dữ liệu (Saving indicator)')
    para(doc, 'Để ứng dụng đạt hiệu năng cao và loại bỏ độ trễ mạng (network latency), TerriMap áp dụng kỹ thuật Cập nhật lạc quan (Optimistic Update) kết hợp với cờ khóa luồng (`saving` flag):')
    bullet(doc, 'Cập nhật lạc quan: Khi người dùng kéo thả thay đổi khu vực của một zone trên bản đồ, Zustand store ngay lập tức cập nhật mảng `assignments` cục bộ trên bộ nhớ RAM. Bản đồ Leaflet re-render ngay lập tức trong vòng 16ms, tạo cảm giác mượt mà 60 FPS.')
    bullet(doc, 'Cờ saving indicator: Ngay khi ghi cục bộ xong, store bật cờ `saving = true` và gửi yêu cầu cập nhật bất đồng bộ lên Supabase Database. Cờ này thông báo cho giao diện hiển thị thông báo "Đang đồng bộ..." và tạm thời vô hiệu hóa (disable) các nút tương tác nhạy cảm để tránh tình trạng Race Condition (xung đột ghi đè dữ liệu).')
    bullet(doc, 'Phục hồi khi lỗi: Nếu kết nối mạng lỗi và việc cập nhật lên database thất bại, store sẽ tự động khôi phục (rollback) trạng thái assignments về giá trị cũ trước khi thay đổi, đồng thời hiển thị thông báo lỗi lên màn hình để người dùng biết.')

    h2(doc, '3.3. Thuật toán Offline-First LocalStorage và đồng bộ Supabase')
    para(doc, 'TerriMap tích hợp cơ chế Offline-First thông minh. Khi ứng dụng mất mạng (được phát hiện qua hàm `navigator.onLine` hoặc lỗi kết nối Supabase):')
    bullet(doc, 'Dữ liệu thay đổi được lưu trữ trực tiếp vào LocalStorage với khóa đặc trưng theo Project ID nhằm tránh rò rỉ chéo dữ liệu giữa các dự án: `terrimap_offline_assignments_{projectId}`.')
    bullet(doc, 'Khi trình duyệt phát hiện có kết nối mạng trở lại (sự kiện `online`), hệ thống sẽ kích hoạt hàm đồng bộ tự động `syncOfflineData()`. Hàm này sẽ đọc dữ liệu từ LocalStorage, thực hiện gom cụm và thực thi câu lệnh `upsert` đa dòng (bulk upsert) lên Supabase để ghi nhận mọi thay đổi một cách nhanh gọn.')

    # ===================== PHẦN 4 =====================
    doc.add_page_break()
    h1(doc, '4. PHÂN HỆ HIỂN THỊ UI/UX VÀ BẢN ĐỒ TƯƠNG TÁC LEAFLET (L4 - PRESENTATION)')
    
    h2(doc, '4.1. Thư viện hiển thị được sử dụng')
    para(doc, 'Giao diện người dùng được thiết kế tỉ mỉ, bóng bẩy và mang lại trải nghiệm chuyên nghiệp cao nhờ kết hợp các thư viện cao cấp:')
    bullet(doc, 'Leaflet & React-Leaflet: Quản lý lớp bản đồ nền (TileLayer từ OpenStreetMap), vẽ các polygon địa lý của Zone và quản lý các sự kiện chuột (click, hover, popup).')
    bullet(doc, 'Recharts: Vẽ đồ thị dạng ResponsiveContainer dưới chân bản đồ. Đồ thị BarChart hiển thị lượng khách hàng thực tế được phân bổ cho từng Sales Agent, giúp người điều phối (Coordinator) đánh giá trực quan mức độ lệch cân bằng giữa các Agent mà không cần xem bảng số liệu khô khan.')
    bullet(doc, 'Tailwind CSS: Dùng để thiết kế giao diện Glassmorphism hiện đại, bóng mờ, viền mỏng và hỗ trợ tối đa Dark Mode để giảm mỏi mắt cho người dùng khi làm việc ban đêm.')

    h2(doc, '4.2. Quản lý trạng thái UI Store (uiStore.ts) chuyên biệt')
    para(doc, 'Tất cả các trạng thái giao diện thuần túy được lưu trữ tách biệt hoàn toàn khỏi dữ liệu nghiệp vụ để tăng tính module hóa của mã nguồn:')
    add_code_block(doc, [
        "export const useUIStore = create<UIState>((set) => ({",
        "  role: 'coordinator',",
        "  selectedZoneId: null,",
        "  highlightedSalesId: null,",
        "  isAlgorithmRunning: false,",
        "  isMapTransitioning: false,",
        "  theme: 'light',",
        "  locale: 'vi',",
        "  setRole: (role) => set({ role }),",
        "  setSelectedZoneId: (id) => set({ selectedZoneId: id }),",
        "  setHighlightedSalesId: (id) => set({ highlightedSalesId: id }),",
        "  setAlgorithmRunning: (running) => set({ isAlgorithmRunning: running }),",
        "  triggerMapTransition: () => {",
        "    set({ isMapTransitioning: true });",
        "    setTimeout(() => set({ isMapTransitioning: false }), 1000);",
        "  }",
        "}));"
    ])

    h2(doc, '4.3. Thành phần TerritoryMap.tsx và cơ chế MapFlyTo nâng cao')
    para(doc, 'Vì thư viện Leaflet khởi tạo tâm bản đồ một lần duy nhất, việc thay đổi region hay project sẽ làm bản đồ bị lệch và không tự động di chuyển. Để xử lý, TerriMap xây dựng component con `MapFlyTo` lồng trực tiếp bên trong `MapContainer`:')
    add_code_block(doc, [
        "interface MapFlyToProps {",
        "  center: [number, number];",
        "  zoom: number;",
        "}",
        "function MapFlyTo({ center, zoom }: MapFlyToProps) {",
        "  const map = useMap(); // Truy cập trực tiếp instance của bản đồ Leaflet",
        "  useEffect(() => {",
        "    map.flyTo(center, zoom, {",
        "      animate: true,",
        "      duration: 1.5, // Chuyển động mượt mà trong 1.5 giây",
        "      easeLinearity: 0.25",
        "    });",
        "  }, [center, zoom, map]);",
        "  return null;",
        "}"
    ])

    h2(doc, '4.4. Quy tắc tô màu trực quan và hiển thị chỉ báo lỗi địa lý')
    para(doc, 'Hệ thống bản đồ thiết kế 3 cơ chế hiển thị đặc biệt nhằm nâng cao hiệu quả làm việc của người điều phối và giảm thiểu sai sót dữ liệu:')
    bullet(doc, 'Bảng phối màu District: Hệ thống sử dụng lookup bảng màu gồm 12 mã màu tương phản cao (như xanh ngọc, cam đất, tím lavender) định nghĩa tại `district-colors.ts`. Zone nào được gán về district nào sẽ tự động thay đổi màu nền (fillColor) và màu viền (color) theo district đó.')
    bullet(doc, 'Chỉ báo Zone cô lập (Island Zone): Đối với những zone không giáp cạnh với bất kỳ zone nào khác (do bản đồ vẽ lỗi hoặc nằm biệt lập ngoài đảo), bản đồ sẽ hiển thị đường viền nét đứt màu cam (dashed orange border) kèm độ dày lớn để coordinator nhận biết ngay lập tức.')
    bullet(doc, 'Chỉ báo Cụm đứt gãy (Disconnected District): Đây là tính năng rất cao cấp. Khi một district bị chia cắt thành 2 mảnh không liên thông (ví dụ: zone A và zone C thuộc cụm 1 nhưng bị ngăn cách bởi zone B thuộc cụm 2), hệ thống sẽ chạy BFS kiểm tra liên thông. Nếu phát hiện đứt gãy, toàn bộ các zone thuộc cụm đó sẽ được vẽ viền nét đứt màu đỏ (dashed red border) để cảnh báo coordinator phải điều chỉnh lại.')

    h2(doc, '4.5. Phân tích mã nguồn và chức năng của 5 màn hình chính')
    para(doc, 'Dự án chia làm 5 trang chính rõ rệt dưới thư mục pages:')
    bullet(doc, 'LoginPage.tsx: Chứa form xác thực người dùng sạch sẽ, kết nối với Supabase Auth, tự động lưu JWT vào cookie/localStorage.')
    bullet(doc, 'ProjectSelectPage.tsx: Giao diện trực quan để người dùng lựa chọn dự án địa lý làm việc hoặc Admin tạo dự án mới (chọn tỉnh thành, gán tên và mô tả).')
    bullet(doc, 'AdminPage.tsx: Dashboard quản lý dữ liệu nền, chứa danh sách Zones và Agents dưới dạng bảng biểu. Cho phép chỉnh sửa nhanh chỉ số hoạt động của từng zone.')
    bullet(doc, 'CoordinatorPage.tsx: Phân hệ cốt lõi. Chứa thanh sidebar cấu hình giải thuật phân hoạch, bản đồ Leaflet ở trung tâm và đồ thị cân bằng tải Recharts bên dưới. Hỗ trợ chạy các luồng tính toán tối ưu song song.')
    bullet(doc, 'SalesPage.tsx: Màn hình tối giản dành riêng cho Sales Agents, chỉ hiển thị duy nhất vùng bản đồ được phân chia cho agent đó và form gửi phản hồi (Feedback) chất lượng vùng.')

    # ===================== PHẦN 5 =====================
    doc.add_page_break()
    h1(doc, '5. PHÂN TÍCH CHI TIẾT THUẬT TOÁN PHÂN VÙNG VÀ LIÊN QUAN ĐẾN CODE')
    
    h2(doc, '5.1. Lựa chọn hạt nhân tối ưu (Farthest-Point Seeding Heuristic)')
    para(doc, 'Để bắt đầu quá trình phân cụm, hệ thống cần chọn ra m zone hạt nhân (seeds) làm tâm phát triển cho m district. Để tránh việc các hạt nhân nằm quá sát nhau dẫn đến các cụm bị chồng chéo, hệ thống áp dụng thuật toán Heuristic chọn điểm xa nhất (Farthest-Point Seeding).')
    para(doc, 'Độ phức tạp thuật toán: O(m * n) thời gian và O(n) không gian lưu trữ bitmask.')
    para(doc, 'Đoạn mã nguồn chi tiết tại `lib/partition.ts` (hàm `selectFarthestSeeds`):')
    add_code_block(doc, [
        "function selectFarthestSeeds(zones: Zone[], m: number): number[] {",
        "  if (zones.length === 0 || m <= 0) return [];",
        "  if (m >= zones.length) return zones.map((_, i) => i);",
        "  const seeds: number[] = [0]; // Chọn seed đầu tiên là phần tử index 0",
        "  const inSeed = new Uint8Array(zones.length);",
        "  inSeed[0] = 1;",
        "  const minDist = new Float64Array(zones.length).fill(Infinity);",
        "  for (let j = 0; j < zones.length; j++) {",
        "    minDist[j] = haversineDistance(zones[0]!.centroid, zones[j]!.centroid);",
        "  }",
        "  minDist[0] = -1;",
        "  for (let s = 1; s < m; s++) {",
        "    let farthest = -1; let maxD = -1;",
        "    for (let j = 0; j < zones.length; j++) {",
        "      if (inSeed[j]) continue;",
        "      if (minDist[j]! > maxD) { maxD = minDist[j]!; farthest = j; }",
        "    }",
        "    if (farthest === -1) break;",
        "    seeds.push(farthest); inSeed[farthest] = 1;",
        "    for (let j = 0; j < zones.length; j++) {",
        "      if (inSeed[j]) continue;",
        "      const d = haversineDistance(zones[farthest]!.centroid, zones[j]!.centroid);",
        "      if (d < minDist[j]!) minDist[j] = d;",
        "    }",
        "  }",
        "  return seeds;",
        "}"
    ])

    h2(doc, '5.2. Thuật toán Greedy Seed Expansion và chiến lược liên thông Grow-to-Reach')
    para(doc, 'Thuật toán Greedy loang biên bắt đầu từ các hạt nhân đã chọn, mở rộng dần ra xung quanh thông qua BFS. Tại mỗi bước, mỗi district chọn trong danh sách zone láng giềng kề ranh giới (frontier) chưa gán một zone có chỉ số lượng khách hàng lớn nhất để sáp nhập.')
    para(doc, 'Trong trường hợp biên BFS của các cụm bị chặn (do không còn zone kề ranh giới chưa gán), thuật toán sẽ kích hoạt chiến lược đột phá liên thông Grow-to-Reach: Tìm đường đi ngắn nhất trên đồ thị kề từ các zone bị cô lập đến zone đã gán gần nhất, rồi gán toàn bộ các zone trên đường đi đó cho cùng một cụm. Điều này đảm bảo 100% các zone đều được gán và giữ tính liên thông đồ thị.')
    para(doc, 'Đoạn mã nguồn chi tiết của Grow-to-Reach trong hàm `partitionGreedy` tại `lib/partition.ts`:')
    add_code_block(doc, [
        "    if (!progress) {",
        "      const unassignedIdxs: number[] = [];",
        "      for (let i = 0; i < zones.length; i++) {",
        "        if (assignment[i] === -1) unassignedIdxs.push(i);",
        "      }",
        "      for (const startIdx of unassignedIdxs) {",
        "        if (assignment[startIdx] !== -1) continue;",
        "        const pathResult = bfsShortestPathToAssigned(zones, adjMatrix, idToIdx, assignment, startIdx);",
        "        if (pathResult) {",
        "          const { path, targetDistrict } = pathResult;",
        "          for (const idx of path) {",
        "            if (assignment[idx] === -1) {",
        "              assignment[idx] = targetDistrict; unassigned--;",
        "              const zId = zones[idx]!.id;",
        "              for (const neighborId of (adjMatrix[zId] ?? [])) {",
        "                const nIdx = idToIdx.get(neighborId);",
        "                if (nIdx !== undefined && assignment[nIdx] === -1) frontiers[targetDistrict]!.add(nIdx);",
        "              }",
        "            }",
        "          }",
        "          progress = true;",
        "        } else {",
        "          let nearestDistrict = 0; let minD = Infinity;",
        "          for (let d = 0; d < m; d++) {",
        "            const seedIdx = seedIndices[d]!;",
        "            const dist = haversineDistance(zones[startIdx]!.centroid, zones[seedIdx]!.centroid);",
        "            if (dist < minD) { minD = dist; nearestDistrict = d; }",
        "          }",
        "          assignment[startIdx] = nearestDistrict; unassigned--; progress = true;",
        "        }",
        "      }",
        "    }"
    ])

    h2(doc, '5.3. Thuật toán Local Search (2-opt Improvement) kiểm soát biên liên thông')
    para(doc, 'Local Search bắt đầu từ kết quả của Greedy, sau đó lặp qua các zone nằm ở biên giới giữa các district và thử chuyển chúng sang district lân cận. Điều kiện tiên quyết để chấp nhận di chuyển là sự di chuyển đó phải làm giảm cost tổng của hàm mục tiêu, và KHÔNG được phá vỡ tính liên thông địa lý của district nguồn. Thuật toán sử dụng hàm `isDistrictConnected` chạy BFS để kiểm tra liên thông cứng này.')
    para(doc, 'Mã nguồn hàm kiểm tra liên thông địa lý `isDistrictConnected` bằng BFS:')
    add_code_block(doc, [
        "export function isDistrictConnected(",
        "  zones: Zone[], assignment: number[], districtId: number,",
        "  adjMatrix: AdjacencyMatrix, idToIdx: Map<string, number>",
        "): boolean {",
        "  const membersIds: string[] = [];",
        "  for (let i = 0; i < zones.length; i++) {",
        "    if (assignment[i] === districtId) membersIds.push(zones[i]!.id);",
        "  }",
        "  if (membersIds.length <= 1) return true;",
        "  const memberSet = new Set(membersIds);",
        "  const visited = new Set<string>();",
        "  const queue = [membersIds[0]!];",
        "  while (queue.length > 0) {",
        "    const current = queue.pop()!;",
        "    if (visited.has(current)) continue;",
        "    visited.add(current);",
        "    for (const neighborId of (adjMatrix[current] ?? [])) {",
        "      if (memberSet.has(neighborId) && !visited.has(neighborId)) {",
        "        queue.push(neighborId);",
        "      }",
        "    }",
        "  }",
        "  return visited.size === memberSet.size; // Kích thước bằng nhau = cụm liên thông",
        "}"
    ])

    h2(doc, '5.4. Thuật toán Simulated Annealing (Ủ mô phỏng) & Tối ưu hóa phân phối Boltzmann')
    para(doc, 'Simulated Annealing (SA) là một thuật toán Metaheuristic có khả năng nhảy ra khỏi cực trị địa phương (local optima) nhờ cơ chế chấp nhận nghiệm tệ hơn với một xác suất nhỏ giảm dần theo nhiệt độ (nhiệt độ T0 hạ dần theo cooling_rate):')
    add_formula(doc, "P(accept) = e^{-\\frac{\\Delta E}{T}} > random[0, 1)")
    para(doc, 'Hàm cost tích hợp 3 thành phần tối ưu hóa đa mục tiêu:')
    add_formula(doc, "Cost = \\alpha \\times Dispersion (p-center) + \\beta \\times Imbalance (std-dev) + \\gamma \\times FragmentsPenalty")
    para(doc, 'Đoạn mã nguồn chi tiết của vòng lặp SA tại `lib/partition.ts`:')
    add_code_block(doc, [
        "  for (let iter = 0; iter < maxIter && T >= 1; iter++) {",
        "    const boundaryZones: number[] = [];",
        "    // Thu thập các zone nằm ở biên giới các cụm",
        "    for (let i = 0; i < zones.length; i++) {",
        "      const dId = assignment[i]!;",
        "      for (const neighborId of (adjMatrix[zones[i]!.id] ?? [])) {",
        "        const nIdx = idToIdx.get(neighborId);",
        "        if (nIdx !== undefined && assignment[nIdx] !== dId) { boundaryZones.push(i); break; }",
        "      }",
        "    }",
        "    if (boundaryZones.length === 0) break;",
        "    const zoneIdx = boundaryZones[Math.floor(Math.random() * boundaryZones.length)]!;",
        "    const currentDistrict = assignment[zoneIdx]!;",
        "    const neighborDistricts = new Set<number>();",
        "    for (const neighborId of (adjMatrix[zones[zoneIdx]!.id] ?? [])) {",
        "      const nIdx = idToIdx.get(neighborId);",
        "      if (nIdx !== undefined && assignment[nIdx] !== currentDistrict) neighborDistricts.add(assignment[nIdx]!);",
        "    }",
        "    if (neighborDistricts.size === 0) continue;",
        "    const targetDistrict = Array.from(neighborDistricts)[Math.floor(Math.random() * neighborDistricts.size)]!;",
        "    assignment[zoneIdx] = targetDistrict; // Thử hoán đổi",
        "    // Kiểm tra liên thông cứng trước khi tính toán chi phí",
        "    if (!isDistrictConnected(zones, Array.from(assignment), currentDistrict, adjMatrix, idToIdx)) {",
        "      assignment[zoneIdx] = currentDistrict; continue; // reject",
        "    }",
        "    const newCost = computeCost(zones, Array.from(assignment), m, alpha, beta, adjMatrix, balanceWeights, objective);",
        "    const deltaE = newCost - currentCost;",
        "    if (deltaE < 0 || Math.random() < Math.exp(-deltaE / T)) {",
        "      currentCost = newCost;",
        "      if (currentCost < bestCost) { bestCost = currentCost; bestAssignment = new Int32Array(assignment); }",
        "    } else { assignment[zoneIdx] = currentDistrict; }",
        "    T *= cooling; // Hạ nhiệt",
        "  }"
    ])

    h2(doc, '5.5. Song song hóa luồng phụ qua Web Workers giải quyết nghẽn Main Thread')
    para(doc, 'Khi số lượng zone tăng lên và số vòng lặp tối đa là 10,000, Simulated Annealing sẽ chiếm dụng 100% năng lực xử lý của một nhân CPU. Nếu chạy trực tiếp trên luồng chính (Main Thread), nó sẽ khiến trình duyệt bị treo (UI freeze), không thể nhận phản hồi chuột hay re-render bản đồ.')
    para(doc, 'TerriMap giải quyết triệt để vấn đề này bằng cách đưa giải thuật SA vào luồng phụ độc lập thông qua Web Workers (`src/workers/sa-worker.ts`). Luồng phụ này tính toán độc lập và chỉ gửi báo cáo tiến trình (`onProgress`) về luồng chính sau mỗi 100 vòng lặp hoặc khi hoàn thành, đảm bảo giao diện web luôn mượt mà 60 FPS.')

    # ===================== PHẦN 6 =====================
    doc.add_page_break()
    h1(doc, '6. CÁC CÔNG THỨC TOÁN HỌC VÀ HÌNH HỌC CỐT LÕI')
    
    h2(doc, '6.1. Công thức khoảng cách mặt cầu Haversine và chống tràn số thực')
    para(doc, 'Để tính khoảng cách địa lý chính xác giữa tâm (centroids) của các zone trên bề mặt cong của Trái Đất, hệ thống sử dụng công thức lượng giác mặt cầu Haversine thay cho khoảng cách Euclid phẳng:')
    add_formula(doc, "d = 2 \\times R \\times \\arcsin\\left(\\sqrt{\\sin^2\\left(\\frac{\\Delta \\varphi}{2}\\right) + \\cos(\\varphi_1)\\cos(\\varphi_2)\\sin^2\\left(\\frac{\\Delta \\lambda}{2}\right)}\\right)")
    para(doc, 'Trong đó R = 6371 km là bán kính trung bình của Trái Đất theo chuẩn WGS-84. delta_phi và delta_lambda là hiệu số vĩ độ và kinh độ được chuyển đổi về đơn vị radians.')
    para(doc, 'Đoạn mã nguồn hàm `haversineDistance` tại `lib/geometry.ts` tích hợp cơ chế clamp đặc biệt để phòng ngừa sai số làm tròn số thực dấu phẩy động (floating-point noise) làm cho giá trị trong căn lớn hơn 1 hoặc nhỏ hơn 0, gây ra lỗi trả về `NaN`:')
    add_code_block(doc, [
        "export function haversineDistance(a: LatLng, b: LatLng): number {",
        "  const R = 6371; // km",
        "  const dLat = toRad(b.lat - a.lat);",
        "  const dLng = toRad(b.lng - a.lng);",
        "  const sinHalfLat = Math.sin(dLat / 2);",
        "  const sinHalfLng = Math.sin(dLng / 2);",
        "  const h = sinHalfLat * sinHalfLat + Math.cos(toRad(a.lat)) * Math.cos(toRad(b.lat)) * sinHalfLng * sinHalfLng;",
        "  // Clamp [0, 1] để loại bỏ floating-point noise sinh ra NaN",
        "  const c = 2 * Math.asin(Math.min(1, Math.sqrt(Math.max(0, h))));",
        "  const dist = R * c;",
        "  return dist === 0 ? 0 : dist;",
        "}"
    ])

    h2(doc, '6.2. Công thức tâm hình học đa giác Shoelace và phòng vệ đa giác suy biến')
    para(doc, 'Tâm hình học (centroid) của một đa giác (zone) là điểm đại diện cho vị trí địa lý của zone đó. Nó được tính toán dựa trên diện tích có hướng của đa giác thông qua công thức dây giày (Shoelace):')
    add_formula(doc, "A = \\frac{1}{2} \\sum_{i=0}^{n-1} (x_i y_{i+1} - x_{i+1} y_i)")
    add_formula(doc, "C_x = \\frac{1}{6A} \\sum_{i=0}^{n-1} (x_i + x_{i+1})(x_i y_{i+1} - x_{i+1} y_i)")
    add_formula(doc, "C_y = \\frac{1}{6A} \\sum_{i=0}^{n-1} (y_i + y_{i+1})(x_i y_{i+1} - x_{i+1} y_i)")
    para(doc, 'Trong đó tọa độ x và y tương ứng với Kinh độ (Longitude) và Vĩ độ (Latitude) của các đỉnh đa giác. Trường hợp đa giác bị suy biến thành đường thẳng hoặc điểm (lỗi dữ liệu thô đầu vào), diện tích A sẽ bằng 0, dẫn đến lỗi chia cho 0. Tệp tin `lib/geometry.ts` thiết lập cơ chế phòng ngự (Defensive Fallback): nếu diện tích |A| < 1e-12, hệ thống tự động đổi sang tính trung bình cộng số học (Arithmetic Mean) tọa độ các đỉnh để làm tâm, loại bỏ lỗi crash ứng dụng.')

    # ===================== PHẦN 7 =====================
    doc.add_page_break()
    h1(doc, '7. THIẾT KẾ CƠ SỞ DỮ LIỆU VÀ CƠ CHẾ BẢO MẬT PHÂN QUYỀN (SUPABASE RLS & RBAC)')
    
    h2(doc, '7.1. Lược đồ quan hệ thực thể cơ sở dữ liệu (Database Schema)')
    para(doc, 'Cơ sở dữ liệu của dự án sử dụng hệ quản trị PostgreSQL quản lý bởi Supabase. Dưới đây là cấu trúc bảng chi tiết kèm theo các khóa chính, khóa ngoại kết nối:')
    bullet(doc, 'profiles (id UUID PK, email Text, full_name Text, avatar_url Text, created_at Timestamp).')
    bullet(doc, 'projects (id Text PK, name Text, description Text, owner_id UUID FK -> profiles(id), created_at Timestamp).')
    bullet(doc, 'project_members (id Text PK, project_id Text FK -> projects(id), user_id UUID FK -> profiles(id), role Text (admin/coordinator/sales), created_at).')
    bullet(doc, 'regions (id Text PK, project_id Text FK, name Text, center_lat Float, center_lng Float, zoom Integer).')
    bullet(doc, 'zones (id Text PK, region_id Text FK, project_id Text FK, name Text, boundary GeoJSON, centroid LatLng, created_at).')
    bullet(doc, 'activities (id Text PK, zone_id Text FK -> zones(id), type Text (CUSTOMER/ORDER/REVENUE), value Float).')
    bullet(doc, 'assignments (zone_id Text PK FK, project_id Text FK, sales_agent_id Text FK, assigned_at Timestamp).')

    h2(doc, '7.2. Chính sách Row Level Security (RLS) bảo mật đa dự án')
    para(doc, 'Để đảm bảo người dùng của dự án A không thể đọc trộm hay sửa đổi dữ liệu của dự án B, Supabase kích hoạt Row Level Security (RLS) ở mức nhân PostgreSQL. Mọi câu lệnh SQL truy vấn từ client đều phải vượt qua chính sách bảo mật sau:')
    add_code_block(doc, [
        "-- Bật RLS cho bảng zones",
        "ALTER TABLE zones ENABLE ROW LEVEL SECURITY;",
        "",
        "-- Tạo chính sách cho phép đọc: chỉ thành viên của project mới được xem các zone tương ứng",
        "CREATE POLICY \"Users can view zones of their projects\" ON zones",
        "FOR SELECT USING (",
        "  EXISTS (",
        "    SELECT 1 FROM project_members",
        "    WHERE project_members.project_id = zones.project_id",
        "      AND project_members.user_id = auth.uid()",
        "  )",
        ");",
        "",
        "-- Tạo chính sách cho phép sửa đổi: chỉ Coordinator hoặc Admin của project mới được sửa zone",
        "CREATE POLICY \"Coordinators can modify zones\" ON zones",
        "FOR ALL USING (",
        "  EXISTS (",
        "    SELECT 1 FROM project_members",
        "    WHERE project_members.project_id = zones.project_id",
        "      AND project_members.user_id = auth.uid()",
        "      AND project_members.role IN ('admin', 'coordinator')",
        "  )",
        ");"
    ])

    h2(doc, '7.3. Triggers tự động hóa đồng bộ tài khoản người dùng')
    para(doc, 'Hệ thống thiết lập một Database Trigger tự động tạo bản ghi trong bảng `profiles` ngay khi có tài khoản mới đăng ký thành công thông qua Supabase Auth:')
    add_code_block(doc, [
        "CREATE OR REPLACE FUNCTION public.handle_new_user()",
        "RETURNS trigger AS $$",
        "BEGIN",
        "  INSERT INTO public.profiles (id, email, full_name, avatar_url)",
        "  VALUES (",
        "    new.id,",
        "    new.email,",
        "    coalesce(new.raw_user_meta_data->>'full_name', 'User'),",
        "    new.raw_user_meta_data->>'avatar_url'",
        "  );",
        "  RETURN new;",
        "END;",
        "$$ LANGUAGE plpgsql SECURITY DEFINER;",
        "",
        "CREATE TRIGGER on_auth_user_created",
        "  AFTER INSERT ON auth.users",
        "  FOR EACH ROW EXECUTE FUNCTION public.handle_new_user();"
    ])

    # ===================== PHẦN 8 =====================
    doc.add_page_break()
    h1(doc, '8. BỘ 30 CÂU HỎI PHẢN BIỆN CHUYÊN SÂU VÀ LỜI GIẢI THÍCH HỌC THUẬT')
    para(doc, 'Dưới đây là tập hợp 30 câu hỏi phản biện cực kỳ hóc búa kèm câu trả lời chi tiết nhất, được thiết kế để sinh viên IT bảo vệ xuất sắc đồ án trước Hội đồng chấm thi:')
    
    questions = [
        ("Câu hỏi 1: Tại sao bài toán phân chia vùng thương mại này lại thuộc nhóm NP-hard?", 
         "Trả lời: Bài toán này thực chất là sự kết hợp của bài toán Phân cụm có ràng buộc dung lượng (Capacitated Clustering Problem - CCP) và bài toán Phân hoạch đồ thị liên thông (Connected Graph Partitioning). Cả hai bài toán này đều đã được chứng minh là NP-hard trong lý thuyết khoa học máy tính. Do số lượng phương án phân hoạch tăng theo hàm mũ (số Stirling loại hai), không có thuật toán thời gian đa thức nào có thể tìm ra nghiệm tối ưu toàn cục cho mọi trường hợp. Do đó chúng ta bắt buộc phải sử dụng các thuật toán xấp xỉ Heuristic và Metaheuristic."),
        
        ("Câu hỏi 2: Sự khác biệt về mặt bản chất giữa thuật toán Greedy Seed Expansion và Local Search là gì?",
         "Trả lời: Greedy Seed Expansion là thuật toán kiến tạo (constructive algorithm). Nó bắt đầu từ tập hạt nhân rỗng và từng bước xây dựng giải pháp bằng cách loang BFS địa lý để gán các zone chưa phân chia. Khi tất cả các zone được gán, thuật toán dừng lại và không quay đầu tối ưu lại. Ngược lại, Local Search là thuật toán cải tiến (improvement algorithm). Nó bắt đầu từ một giải pháp đã hoàn chỉnh (kết quả của Greedy) và liên tục thử thực hiện các phép hoán đổi cục bộ (swap) ở biên để tìm cách giảm chi phí hàm mục tiêu, cải thiện chất lượng nghiệm hiện có."),
         
        ("Câu hỏi 3: Tại sao Simulated Annealing lại tốt hơn Local Search trong bài toán này?",
         "Trả lời: Local Search hoạt động theo cơ chế leo đồi đơn giản (greedy hill-climbing), nghĩa là nó chỉ chấp nhận các phép hoán đổi làm giảm cost ngay lập tức. Điều này khiến nó dễ dàng bị kẹt ở các cực trị cục bộ (local optima) và không thể tìm được nghiệm tốt hơn. Simulated Annealing giải quyết điều này bằng cách cho phép chấp nhận các phép hoán đổi xấu hơn với một xác suất Boltzmann nhỏ giảm dần theo nhiệt độ. Cơ chế này giúp giải thuật có khả năng 'nhảy' ra khỏi các thung lũng cục bộ để tìm kiếm các vùng không gian nghiệm tốt hơn toàn cục."),
         
        ("Câu hỏi 4: Cờ 'initialized' trong Zustand Store giải quyết bottleneck gì của React?",
         "Trả lời: Trong React, khi người dùng chuyển đổi qua lại giữa các màn hình (ví dụ: chuyển từ Admin sang Coordinator), các React Components sẽ bị unmount và hủy hoàn toàn trạng thái cục bộ. Nếu không có store Zustand, mỗi lần chuyển trang React sẽ phải gọi API Supabase tải lại dữ liệu từ đầu. Cờ `initialized` giúp store nhận biết dữ liệu đã có trên RAM của trình duyệt rồi để bỏ qua việc fetch lại, giúp chuyển đổi trang tức thời và giảm tải hàng ngàn truy vấn không cần thiết lên Database."),
         
        ("Câu hỏi 5: Giải thích chi tiết chiến lược 'Grow-to-Reach' xử lý tắc nghẽn biên BFS của bạn?",
         "Trả lời: Trong thuật toán Greedy loang biên BFS, khi ranh giới của các cụm đã chạm vào nhau, có khả năng vẫn còn một số zone chưa được gán nằm xa hoặc bị kẹt mà biên BFS không loang tới được (tắc nghẽn). Thay vì gán bừa bãi phá vỡ tính liên thông, hệ thống sẽ chạy BFS tìm đường đi ngắn nhất trên đồ thị kề từ zone bị kẹt này đến zone đã gán gần nhất. Sau đó, nó gán toàn bộ các zone trung gian dọc theo con đường này cho cụm đích, tạo ra một 'hành lang liên thông' vững chắc nối zone cô lập về cụm."),
         
        ("Câu hỏi 6: Tại sao bạn lại clamp giá trị trong hàm tính khoảng cách Haversine?",
         "Trả lời: Trong tính toán số thực dấu phẩy động (IEEE 754), các phép nhân lượng giác có thể tích lũy sai số làm tròn siêu nhỏ khiến giá trị h tính được lớn hơn 1 một chút (ví dụ: 1.0000000000000002). Khi đưa giá trị này vào hàm căn `Math.sqrt` và hàm lượng giác ngược `Math.asin`, JavaScript sẽ trả về kết quả `NaN` (Not a Number) và gây crash thuật toán. Việc dùng `Math.min(1, Math.sqrt(Math.max(0, h)))` đảm bảo đầu vào của hàm `Math.asin` luôn nằm trong đoạn an toàn [-1, 1], ngăn chặn hoàn toàn lỗi NaN."),
         
        ("Câu hỏi 7: Phân tích độ phức tạp thời gian (Time Complexity) của thuật toán Greedy Seed Expansion?",
         "Trả lời: Thuật toán Greedy gồm 2 pha: (1) Chọn m hạt nhân bằng farthest-point heuristic mất O(m * n). (2) Loang BFS mở rộng biên giới. Ở mỗi vòng lặp, ta duyệt qua biên frontiers của m cụm để chọn candidate tốt nhất. Số lần gán zone là (n - m). Với mỗi lần gán, ta cập nhật láng giềng k kề ranh giới. Độ phức tạp tổng quát của pha loang là O(n * log(k) * m) trong trường hợp thông thường, hoặc tối đa O(n^2) trong trường hợp xấu nhất đồ thị kề dày đặc. Đây là tốc độ cực nhanh, thực thi dưới 1ms cho quy mô hàng trăm zone."),
         
        ("Câu hỏi 8: Làm thế nào thuật toán isDistrictConnected chạy BFS kiểm tra liên thông có thể ngăn chặn lỗi đứt gãy địa lý?",
         "Trả lời: Khi thuật toán SA hoặc Local Search thử hoán đổi zone X từ cụm A sang cụm B, có nguy cơ zone X chính là chiếc cầu nối duy nhất giữ các phần của cụm A liên thông. Để kiểm tra, ta tạm thời gán zone X cho cụm B, sau đó chạy BFS bắt đầu từ một zone bất kỳ của cụm A. BFS sẽ loang qua ma trận kề để đếm số zone của cụm A có thể đi tới được. Nếu số lượng zone duyệt được bằng đúng số lượng zone thực tế của cụm A thì cụm A vẫn liên thông (chấp nhận hoán đổi). Nếu nhỏ hơn, cụm A đã bị đứt gãy thành hai mảnh tách rời (lập tức reject phép hoán đổi)."),
         
        ("Câu hỏi 9: Tại sao bạn không chạy Simulated Annealing trên Main Thread của trình duyệt?",
         "Trả lời: Trình duyệt JavaScript chạy đơn luồng. Main Thread chịu trách nhiệm cho cả việc re-render giao diện, xử lý sự kiện click chuột và chạy code JS. Thuật toán SA có tới 10,000 vòng lặp, mỗi vòng lặp chạy BFS kiểm tra liên thông và tính khoảng cách Haversine. Nếu chạy SA trực tiếp trên Main Thread, nó sẽ block hoàn toàn luồng này trong khoảng 200ms - 1000ms. Trong thời gian này, trình duyệt sẽ bị đơ hoàn toàn (UI freeze), người dùng không thể cuộn trang hay bấm nút. Do đó, đưa SA vào Web Worker chạy nền là bắt buộc để giữ UI mượt mà 60 FPS."),
         
        ("Câu hỏi 10: RLS (Row Level Security) của Supabase hoạt động ở tầng nào và bảo mật dự án thế nào?",
         "Trả lời: RLS hoạt động ở tầng nhân Database của PostgreSQL, chứ không phải ở tầng code Node.js/React. Khi client gửi một truy vấn đọc bảng `zones`, Supabase tự động đính kèm JSON Web Token (JWT) của user hiện tại. PostgreSQL sẽ kiểm tra xem ID người dùng này có thuộc bảng `project_members` của dự án đó hay không thông qua chính sách đã tạo. Nếu không, câu lệnh SQL sẽ tự động lọc bỏ hàng dữ liệu đó. Điều này đảm bảo an toàn tuyệt đối, dù hacker có can thiệp sửa code JavaScript ở client cũng không thể đọc trộm dữ liệu của dự án khác."),
         
        ("Câu hỏi 11: Làm thế nào ứng dụng của bạn xử lý lỗi chia cho 0 trong Shoelace Centroid?",
         "Trả lời: Khi tính tâm hình học đa giác (centroid) bằng công thức Shoelace, ta chia mô-men tĩnh cho diện tích 6A. Nếu đa giác bị suy biến (ví dụ: người dùng vẽ lỗi làm các điểm thẳng hàng, đa giác có diện tích A = 0), phép chia cho 6A sẽ tạo ra lỗi chia cho 0, sinh ra tọa độ tâm là Infinity hoặc NaN. Để phòng vệ, hệ thống kiểm tra nếu diện tích |A| < 1e-12, nó sẽ tự động bỏ qua phép tính Shoelace và chuyển sang dùng trung bình cộng đơn giản (Arithmetic Mean) tọa độ của các đỉnh làm tâm, đảm bảo đầu ra luôn là số thực hợp lệ."),
         
        ("Câu hỏi 12: Thiết kế Offline-First của bạn lưu trữ dữ liệu ở đâu và làm sao tránh ghi đè chéo dữ liệu dự án?",
         "Trả lời: Khi mất kết nối mạng, TerriMap ghi dữ liệu tạm thời vào LocalStorage của trình duyệt. Để tránh xung đột ghi đè chéo dữ liệu giữa các dự án khác nhau (ví dụ: người dùng đang mở dự án Hà Nội sau đó chuyển sang dự án Huế), khóa LocalStorage được phân tách (scoped) rõ ràng bằng cách chèn trực tiếp projectId vào tên key: `terrimap_offline_assignments_{projectId}`. Khi có mạng trở lại, hệ thống chỉ đẩy đúng dữ liệu của projectId tương ứng lên DB."),
         
        ("Câu hỏi 13: Tại sao bạn lại sử dụng Int32Array và Float64Array trong giải thuật partition thay vì Array thông thường?",
         "Trả lời: Trong JavaScript, mảng thông thường (Array) là mảng động có thể chứa bất kỳ kiểu dữ liệu nào, điều này khiến trình biên dịch V8 của Chrome không thể tối ưu hóa bộ nhớ và tốc độ truy cập. Bằng cách sử dụng mảng định kiểu (TypedArrays) như `Int32Array` (lưu assignment index) và `Float64Array` (lưu khoảng cách thực), dữ liệu được lưu trữ liên tục trên các khối nhớ byte. Điều này tăng tốc độ truy cập RAM lên nhiều lần và tối ưu hóa bộ nhớ đệm CPU cache khi chạy hàng triệu phép so sánh trong Simulated Annealing."),
         
        ("Câu hỏi 14: K-Means thuần túy gặp nhược điểm gì khi giải quyết bài toán thiết kế vùng thương mại?",
         "Trả lời: K-Means là thuật toán phân cụm không giám sát dựa trên khoảng cách hình học thuần túy. Nó có 2 điểm yếu chí mạng trong bài toán này: (1) Nó không có cơ chế ràng buộc cân bằng tải trọng (ví dụ: số khách hàng trong các cụm phải bằng nhau), dẫn đến cụm thì quá tải, cụm thì quá ít khách. (2) K-Means không đảm bảo tính liên thông địa lý cứng của các cụm trên địa hình phức tạp hoặc có vật cản. Do đó, TerriMap chỉ dùng Greedy loang BFS để tạo nghiệm ban đầu và dùng Simulated Annealing để tối ưu hóa đa mục tiêu hàm cost."),
         
        ("Câu hỏi 15: Phân tích độ phức tạp không gian (Space Complexity) của thuật toán Simulated Annealing?",
         "Trả lời: Thuật toán Simulated Annealing của TerriMap cực kỳ tối ưu về bộ nhớ. Bộ nhớ lưu trữ ma trận kề tiếp giáp địa lý có kích thước O(n * k) với k là số lân cận trung bình. Mảng assignment lưu trữ kết quả gán chỉ tốn O(n) không gian dưới dạng `Int32Array`. Thuật toán thực hiện hoán đổi tại chỗ (in-place swap) và chỉ sao chép một mảng `bestAssignment` có kích thước O(n) khi tìm được nghiệm tốt hơn. Do đó, độ phức tạp không gian chỉ là O(n), hoàn toàn hoạt động nhẹ nhàng trên các thiết bị di động cấu hình yếu."),
         
        ("Câu hỏi 16: Trình bày luồng hoạt động chi tiết của component MapFlyTo trong TerritoryMap.tsx?",
         "Trả lời: Leaflet chỉ khởi tạo tâm bản đồ một lần khi mount component MapContainer. Khi người dùng click chọn một Region mới, tọa độ tâm mới được truyền qua Props của TerritoryMap. Để map tự động di chuyển, component con `MapFlyTo` sử dụng hook `useMap()` để lấy ra thực thể bản đồ Leaflet đang chạy dưới nền. Sau đó, một React `useEffect` sẽ lắng nghe sự thay đổi của cặp tọa độ center [lat, lng]. Khi tọa độ thay đổi, nó sẽ gọi hàm `map.flyTo(center, zoom, { duration: 1.5 })` để di chuyển bản đồ lướt mượt mà đến vùng mới."),
         
        ("Câu hỏi 17: Làm thế nào bạn phát hiện ra một Zone là cô lập (Island Zone) để vẽ viền nét đứt màu cam?",
         "Trả lời: Trong quá trình xây dựng ma trận kề địa lý (Adjacency Matrix), hệ thống duyệt qua tất cả các cặp zone và tính khoảng cách Haversine. Nếu một zone X có khoảng cách đến tất cả các zone khác lớn hơn ngưỡng tiếp giáp (15km), nghĩa là danh sách neighbor của zone X trong ma trận kề là rỗng (empty array). Khi vẽ bản đồ, TerritoryMap sẽ kiểm tra nếu `adjMatrix[zone.id]?.length === 0` thì sẽ áp dụng style riêng biệt với nét vẽ viền đứt quãng nét dày màu cam để cảnh báo."),
         
        ("Câu hỏi 18: Sự khác biệt lớn nhất giữa hàm mục tiêu p-center và p-median trong tùy chọn giải thuật là gì?",
         "Trả lời: (1) p-center (mặc định) hướng tới việc giảm thiểu khoảng cách lớn nhất từ tâm cụm đến bất kỳ zone thành viên nào (minimize max diameter). Nó giúp tạo ra các cụm tròn trịa, gọn gàng nhất và tránh việc một zone ở quá xa bị gom vào cụm. (2) p-median hướng tới việc giảm thiểu tổng khoảng cách của tất cả các zone thành viên đến tâm cụm. Nó phù hợp khi tối ưu hóa tổng chi phí di chuyển trung bình của nhân viên."),
         
        ("Câu hỏi 19: Bạn viết unit tests cho giải thuật ở L1 như thế nào để đảm bảo tính độc lập?",
         "Trả lời: Do tầng L1 (Library) được thiết kế là các Pure Functions, không phụ thuộc vào trạng thái React hay Database, việc viết unit test cực kỳ đơn giản và trực quan. Chúng tôi sử dụng thư viện Vitest, chuẩn bị dữ liệu đầu vào tĩnh dưới dạng mock các đối tượng Zone[], số cụm m và gọi hàm trực tiếp. Bài test kiểm tra các kết quả đầu ra như: tất cả các zone có được gán không, mã cụm có nằm trong đoạn [0, m-1] không, và chạy thử giải thuật với tham số sai để kiểm tra xem nó có ném ra các ngoại lệ `PartitionError` tương ứng hay không."),
         
        ("Câu hỏi 20: Tại sao trong hàm computeCost lại có tham số gamma nhân với totalFragments?",
         "Trả lời: Trong quá trình Simulated Annealing hoán đổi ngẫu nhiên các zone biên, có thể xảy ra trường hợp một hoán đổi tạm thời tạo ra một cụm bị đứt rời địa lý (disconnected). Để loại bỏ hoàn toàn các nghiệm rác này khỏi không gian nghiệm, hàm cost áp dụng kỹ thuật hàm phạt (Penalty Function). Tham số gamma được gán giá trị rất lớn (gamma = 50). Khi cụm bị chia cắt làm c mảnh (c > 1), chi phí sẽ bị cộng thêm (c - 1) * 50. Chi phí phạt khổng lồ này khiến thuật toán SA ngay lập tức từ chối nghiệm đó trong bước so sánh cost, hướng giải thuật tìm kiếm các nghiệm liên thông."),
         
        ("Câu hỏi 21: Database Trigger trong Supabase PostgreSQL của bạn giải quyết vấn đề gì của kiến trúc?",
         "Trả lời: Khi người dùng đăng ký tài khoản mới qua Supabase Auth, thông tin tài khoản được lưu trong bảng hệ thống ẩn `auth.users` của Supabase mà client React không thể truy cập trực tiếp vì lý do bảo mật. Database Trigger giúp tự động bắt sự kiện insert của bảng `auth.users` và ghi đè thông tin tương ứng vào bảng công khai `public.profiles` của chúng ta, giúp đồng bộ hóa tức thời thông tin avatar, email, họ tên người dùng để hiển thị lên UI."),
         
        ("Câu hỏi 22: Hãy giải thích cách thuật toán BFS shortest path hoạt động trong Grow-to-Reach?",
         "Trả lời: BFS shortest path nhận đầu vào là đồ thị kề địa lý và zone xuất phát đang bị cô lập. Nó sử dụng một queue chứa các cặp [zoneIndex, path]. Tại mỗi bước, nó lấy phần tử đầu queue ra, kiểm tra các neighbor của nó. Nếu neighbor đó là một zone đã được gán cho một cụm nào đó, thuật toán lập tức dừng lại và trả về con đường đi ngắn nhất cùng mã cụm đích đó. Nếu neighbor chưa được gán, nó sẽ thêm neighbor và đường đi mới vào queue để tiếp tục loang."),
         
        ("Câu hỏi 23: Làm thế nào bạn tối ưu kích thước gói bundle JS khi build ứng dụng với Leaflet?",
         "Trả lời: Thư viện bản đồ Leaflet có kích thước khá lớn. Để tránh làm chậm thời gian tải trang đầu tiên của người dùng, chúng tôi sử dụng công cụ build Vite hỗ trợ Dynamic Imports (code splitting). Thành phần `TerritoryMap` và thư viện Leaflet được đóng gói riêng thành một chunk JS độc lập và chỉ được tải bất đồng bộ (lazy load) khi người dùng đăng nhập thành công và truy cập vào màn hình làm việc của Admin hoặc Coordinator."),
         
        ("Câu hỏi 24: Tại sao trong tệp db.ts của L2 bạn lại sử dụng bulk insert thay vì vòng lặp insert từng dòng?",
         "Trả lời: Mỗi lần gửi một câu lệnh SQL insert lên Database đám mây, ứng dụng phải tốn thời gian thiết lập kết nối mạng và độ trễ khứ hồi (Round-Trip Time - RTT) mất khoảng 50ms - 100ms. Nếu có 100 assignments và ta dùng vòng lặp để insert từng dòng, tổng thời gian ghi sẽ mất tới 10 giây và làm nghẽn kết nối DB. Bằng cách gom tất cả các assignments thành một mảng và gọi một câu lệnh bulk upsert duy nhất `supabase.from('assignments').upsert(data)`, toàn bộ quá trình ghi chỉ diễn ra trong vòng dưới 100ms."),
         
        ("Câu hỏi 25: Phân tích cơ chế đồng bộ xung đột (Conflict Resolution) của bạn khi đồng bộ LocalStorage lên Supabase?",
         "Trả lời: Khi đồng bộ dữ liệu assignments từ LocalStorage lên Supabase PostgreSQL, hệ thống sử dụng cú pháp `ON CONFLICT (zone_id) DO UPDATE`. Nghĩa là nếu zone đó đã được một người điều phối khác phân chia trên database trong thời gian ta mất mạng, database sẽ ưu tiên cập nhật đè giá trị assignments mới nhất của ta lên (Last-Write-Wins), đảm bảo không phát sinh lỗi trùng khóa chính (Primary Key Violation) làm crash tiến trình đồng bộ."),
         
        ("Câu hỏi 26: Tại sao lại nói thiết kế của L1 (Library) hoàn toàn tuân thủ nguyên tắc Pure Functions?",
         "Trả lời: Vì các hàm trong L1 như `haversineDistance` hay `partitionGreedy` chỉ nhận tham số đầu vào qua các tham số tường minh (arguments), thực hiện tính toán trên bộ nhớ và trả về kết quả qua lệnh `return`. Chúng hoàn toàn không đọc/ghi biến toàn cục nào ngoài hàm, không thực hiện các tác vụ I/O (không gọi mạng, không truy vấn DB), và không thay đổi dữ liệu gốc truyền vào (chỉ tạo bản sao). Điều này đảm bảo tính dự đoán tuyệt đối của chương trình và giúp việc viết các bài kiểm thử tự động đạt độ bao phủ cao."),
         
        ("Câu hỏi 27: Làm thế nào để thay đổi thuật toán phân cụm mặc định trong tương lai mà không cần viết lại giao diện UI?",
         "Trả lời: Nhờ áp dụng mẫu thiết kế Factory Pattern tại L1 (`getPartitionFn`), toàn bộ các thuật toán phân vùng đều tuân thủ một giao thức chữ ký chung (Common Interface) là `PartitionFn`. Nếu trong tương lai ta muốn thêm thuật toán mới (ví dụ: Genetic Algorithm), ta chỉ cần viết thuật toán đó tuân thủ đúng kiểu dữ liệu đầu vào/đầu ra của `PartitionFn` và đăng ký nó vào hàm factory `getPartitionFn`. Giao diện React ở L4 chỉ gọi qua Factory nên hoàn toàn không cần chỉnh sửa bất kỳ dòng code UI nào."),
         
        ("Câu hỏi 28: std-dev (Độ lệch chuẩn) trong hàm tính imbalance giúp đánh giá tải trọng thế nào?",
         "Trả lời: Để đo lường mức độ mất cân bằng tải trọng (ví dụ: lượng khách hàng) giữa m cụm, hệ thống tính giá trị trung bình cộng lượng khách hàng của m cụm. Sau đó, nó tính phương sai (variance) bằng cách lấy trung bình bình phương độ lệch của từng cụm so với mức trung bình. Độ lệch chuẩn (std-dev) là căn bậc hai của phương sai. Trị số độ lệch chuẩn càng nhỏ chứng tỏ sự phân bổ khách hàng giữa các cụm càng đều nhau, độ lệch chuẩn bằng 0 nghĩa là tất cả các cụm có lượng khách bằng khít nhau."),
         
        ("Câu hỏi 29: Làm thế nào bạn validate dữ liệu đa giác GeoJSON ở tầng L0 để tránh lỗi dựng hình bản đồ?",
         "Trả lời: Tại file `lib/validator.ts` thuộc L1, hệ thống thực hiện kiểm thử phòng ngự (Defensive Validation) đối với dữ liệu GeoJSON truyền vào. Hàm validation kiểm tra các ràng buộc: cấu trúc GeoJSON phải đúng định dạng Coordinate[][], số lượng đỉnh của đa giác phải lớn hơn hoặc bằng 3 để tạo thành đa giác kín, tọa độ các đỉnh phải là số thực hợp lệ và nằm trong giới hạn địa lý vĩ độ [-90, 90] và kinh độ [-180, 180]. Nếu vi phạm, hệ thống sẽ chặn ngay từ cổng vào và ném ra lỗi tường minh."),
         
        ("Câu hỏi 30: Đồ thị Recharts so sánh tải trọng giúp Coordinator đưa ra quyết định gì?",
         "Trả lời: Biểu đồ cột của Recharts hiển thị trực quan tổng giá trị tải trọng (ví dụ: số khách hàng) của từng cụm phân chia cạnh nhau. Nếu Coordinator nhìn thấy một cột quá cao và các cột khác quá thấp (mất cân bằng tải), họ có thể ngay lập tức sử dụng chuột click vào cụm quá tải trên bản đồ Leaflet, xem các zone biên và kéo thả thủ công sang cụm lân cận để san sẻ tải trọng, giúp quy hoạch vùng bán hàng tối ưu và thực tế hơn.")
    ]
    
    for q_title, q_ans in questions:
        h2(doc, q_title)
        p_ans = doc.add_paragraph()
        p_ans.paragraph_format.space_after = Pt(6)
        p_ans.paragraph_format.line_spacing = 1.2
        
        run_lbl = p_ans.add_run("Trả lời: ")
        run_lbl.bold = True
        run_lbl.font.name = 'Arial'
        run_lbl.font.size = Pt(10.5)
        run_lbl.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        
        run_txt = p_ans.add_run(q_ans)
        run_txt.font.name = 'Arial'
        run_txt.font.size = Pt(10.5)
        run_txt.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ===================== PHẦN 9 =====================
    doc.add_page_break()
    h1(doc, '9. ĐÁNH GIÁ CHẤT LƯỢNG HỆ THỐNG VÀ ĐỊNH HƯỚNG MỞ RỘNG TƯƠNG LAI')
    para(doc, 'Hệ thống TerriMap Enterprise (v3.5) đã giải quyết trọn vẹn và tối ưu các yêu cầu kỹ thuật và nghiệp vụ phức tạp của bài toán phân chia vùng thương mại thực tế. Việc áp dụng kiến trúc 5 lớp Clean Architecture giúp mã nguồn của dự án đạt độ module hóa cực cao, dễ bảo trì và mở rộng độc lập. Luồng tính toán giải thuật SA kết hợp Web Workers đã mang lại trải nghiệm tương tác mượt mà tuyệt đối cho người dùng.')
    para(doc, 'Định hướng nghiên cứu và phát triển tiếp theo của dự án bao gồm:')
    bullet(doc, 'Tích hợp thêm các ràng buộc mềm từ người điều phối (Soft Constraints) như: gán cứng một số zone bắt buộc phải do Sales Agent A phụ trách vì lý do quan hệ khách hàng quen thuộc.')
    bullet(doc, 'Phát triển giải thuật di truyền đa mục tiêu (Multi-Objective Genetic Algorithm NSGA-II) để tự động xuất ra tập hợp các phương án tối ưu Pareto cho người dùng lựa chọn.')
    bullet(doc, 'Hỗ trợ tính năng cộng tác thời gian thực (Real-time Collaboration) qua WebSockets, cho phép nhiều Coordinator cùng thiết kế và thảo luận trên một bản đồ chung.')
    
    # Save v3
    path = os.path.join(OUTPUT_DIR, 'TerriMap_BaoCaoKyThuat_v3.docx')
    doc.save(path)
    print(f'[OK] Saved technical report to: {path}')
    return path

# ============================================================
# GENERATION: FILE 2 - DEMO SCRIPT
# ============================================================

def build_demo():
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(120)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run('KỊCH BẢN TRÌNH BÀY VÀ DEMO LIVE DỰ ÁN TERRIMAP')
    run_title.font.size = Pt(18)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(200)
    run_sub = p_sub.add_run('Kịch bản thuyết trình chi tiết từng bước kèm lời thoại học thuật và chỉ số benchmark\n(Bản tiếng Việt có dấu hoàn chỉnh)')
    run_sub.font.size = Pt(12)
    run_sub.italic = True
    run_sub.font.color.rgb = RGBColor(0x4A, 0x77, 0x9D)
    
    doc.add_page_break()
    
    h1(doc, 'KỊCH BẢN BẢO VỆ & DEMO LIVE (THỜI LƯỢNG ~25 PHÚT)')
    
    # Phase 1
    h2(doc, 'Phần 1: Giới thiệu bài toán và Đặt vấn đề (Thời lượng: 3 phút)')
    h3(doc, '💬 Lời thoại thuyết trình:')
    para(doc, '"Kính thưa Hội đồng, trong các doanh nghiệp phân phối hàng tiêu dùng hoặc dịch vụ, việc phân chia thị trường cho nhân viên kinh doanh thường được thực hiện thủ công dựa trên cảm tính của người quản lý. '
              'Điều này dẫn đến tình trạng bất công bằng: người thì phụ trách khu vực quá đông khách hàng, người thì phải di chuyển quãng đường quá xa, làm giảm hiệu quả kinh doanh. '
              'TerriMap ra đời nhằm tự động hóa quy trình thiết kế vùng thương mại dựa trên các giải thuật toán học tối ưu, đảm bảo cân bằng công việc và tối ưu hóa quãng đường di chuyển."', italic=True)
              
    h3(doc, '🖥️ Thao tác trực quan trên màn hình:')
    para(doc, 'Mở trình duyệt truy cập ứng dụng. Trỏ chuột vào màn hình bản đồ tương tác hiển thị danh sách các zone tại Hà Nội với đường viền polygon sắc nét.')
    
    # Phase 2
    h2(doc, 'Phần 2: Demo trực quan 3 giải thuật phân vùng (Thời lượng: 12 phút)')
    
    # Greedy
    h3(doc, '1. Thuật toán Greedy Seed Expansion (Thời lượng: 3 phút)')
    para(doc, '💬 Lời thoại thuyết trình:', bold=True)
    para(doc, '"Đầu tiên, chúng tôi demo giải thuật Greedy. Thuật toán sẽ tìm kiếm m seeds hạt nhân xa nhau nhất để bắt đầu loang BFS đồng thời. '
              'Với mỗi bước, cụm sẽ sáp nhập zone lân cận có lượng khách hàng lớn nhất. Điểm mạnh của Greedy là tốc độ chạy cực nhanh, chỉ khoảng 0.5ms cho 20 zones."', italic=True)
    para(doc, '🖥️ Thao tác: Chọn chế độ phân vùng Greedy trên bảng điều khiển, nhập số cụm là 4, nhấn "Chạy phân chia". Bản đồ lập tức tô màu 4 cụm liên thông rõ rệt.', bold=True)
    
    # Local Search
    h3(doc, '2. Thuật toán Local Search (2-opt Improvement) (Thời lượng: 4 phút)')
    para(doc, '💬 Lời thoại thuyết trình:', bold=True)
    para(doc, '"Tiếp theo là Local Search. Thuật toán này cải tiến kết quả từ Greedy bằng cách liên tục duyệt qua các zone giáp ranh đường biên và thử hoán đổi chúng sang cụm lân cận để xem cost hàm mục tiêu có giảm không. '
              'Mỗi bước hoán đổi đều được kiểm tra liên thông nghiêm ngặt bằng BFS. Tốc độ thực thi của giải thuật là khoảng 1.0ms."', italic=True)
    para(doc, '🖥️ Thao tác: Chọn Local Search, nhấn nút chạy. Chỉ số Balance Score cải thiện đáng kể trên màn hình kết quả.', bold=True)
    
    # SA
    h3(doc, '3. Thuật toán Simulated Annealing (Thời lượng: 5 phút)')
    para(doc, '💬 Lời thoại thuyết trình:', bold=True)
    para(doc, '"Cuống cùng là Simulated Annealing (Ủ mô phỏng). Đây là thuật toán mạnh mẽ nhất. SA cho phép chấp nhận các bước dịch chuyển đi lùi (tệ hơn) với xác suất Boltzmann để thoát khỏi cực trị cục bộ. '
              'Để ngăn chặn đơ trình duyệt khi chạy 10,000 iter tính toán, chúng tôi đã sử dụng Web Workers để giải thuật chạy ngầm. Kết quả của SA đạt độ cân bằng tải lên tới 90-99%."', italic=True)
    para(doc, '🖥️ Thao tác: Chọn Simulated Annealing, nhấn chạy. Thấy thanh tiến trình chạy mượt mà và bản đồ cập nhật kết quả phân vùng tối ưu cực kỳ đẹp mắt.', bold=True)
    
    # Benchmarks
    h2(doc, 'Phần 3: So sánh Benchmark thực tế giữa các giải thuật')
    add_styled_table(doc,
        ['Giải thuật', 'Thời gian chạy (ms)', 'Độ cân bằng tải (Balance Score)', 'Đường kính tối đa (Max Diameter)', 'Đảm bảo liên thông địa lý'],
        [
            ['Greedy Seed Expansion', '0.5 ms', 'Trung bình (70 - 85%)', 'Lớn', 'Đảm bảo 100% qua BFS'],
            ['Local Search (2-opt)', '1.0 ms', 'Tốt (80 - 90%)', 'Trung bình', 'Đảm bảo 100% qua check-connected'],
            ['Simulated Annealing (SA)', '2.8 ms', 'Xuất sắc (90 - 99%)', 'Tối ưu nhỏ nhất', 'Đảm bảo 100% qua penalty & check-connected']
        ])
    
    # Phase 3
    h2(doc, 'Phần 4: Tổng kết và Hỏi đáp (Thời lượng: 5 phút)')
    para(doc, 'Sinh viên trình bày tóm tắt các điểm nổi bật của dự án (Clean Architecture, Web Workers, RLS Security, Offline-First) và sẵn sàng nhận các câu hỏi phản biện từ Hội đồng chấm thi.')
    
    # Save v3
    path = os.path.join(OUTPUT_DIR, 'TerriMap_KichBanDemo_v3.docx')
    doc.save(path)
    print(f'[OK] Saved demo script to: {path}')
    return path

# ============================================================
# MAIN EXECUTION
# ============================================================

if __name__ == '__main__':
    print('=== Generating Highly Detailed TerriMap Documents (v3.5) ===')
    p1 = build_report()
    p2 = build_demo()
    print(f'\nSuccess! Word Files generated and saved to Desktop:')
    print(f'  1. {p1}')
    print(f'  2. {p2}')
